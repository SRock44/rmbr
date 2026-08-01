"""Binary document extraction — PDF and DOCX — for `Index.add_files()`.

Both are optional (`pip install rmbr[pdf]` / `rmbr[docx]`), lazily
imported so neither `pypdf` nor `python-docx` is a required rmbr
dependency. This module only turns a binary file into one plain-text
string; `chunk.py` still does all the splitting, same as every other
format `add_files()` handles.

**If the file is found but the extra isn't installed, this raises** —
deliberately, not a silent skip. `add_files()`'s existing skip-unknown-
extensions behavior is for files you plainly didn't mean to index
(images, lockfiles); a `.pdf` sitting right there in the directory you
asked to index is not that — silently leaving it out would produce an
index missing content the caller has every reason to expect is in it.

**Honesty note on quality:** extraction quality is genuinely uneven.
Clean, real text-based PDFs/DOCX extract well. Scanned or image-only
PDFs extract to near-empty text — `pypdf` does not do OCR, deliberately;
that's a much heavier dependency (a real OCR engine, not a parsing
library) for a need most rmbr users won't have. Complex multi-column PDF
layouts can extract out of reading order. If you need OCR, run it
yourself and feed the resulting text to `add_text()` directly — the same
escape hatch as any format rmbr doesn't natively understand.
"""

from __future__ import annotations

from pathlib import Path


def extract_pdf(path: Path) -> str:
    """Plain text from a PDF: every page's extractable text, joined.
    No OCR — a scanned/image-only page extracts to nothing. Requires
    `pypdf` (`pip install rmbr[pdf]`).
    """
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ImportError(
            f"found a PDF ({path}) but pypdf isn't installed - pip install rmbr[pdf], "
            "or pass a pattern= to add_files() that excludes .pdf files."
        ) from exc
    reader = PdfReader(str(path))
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def extract_docx(path: Path) -> str:
    """Plain text from a DOCX: every paragraph, then every table's cells
    (row-major, tab-separated between cells, one line per row). Requires
    `python-docx` (`pip install rmbr[docx]`).
    """
    try:
        import docx
    except ImportError as exc:
        raise ImportError(
            f"found a DOCX ({path}) but python-docx isn't installed - pip install rmbr[docx], "
            "or pass a pattern= to add_files() that excludes .docx files."
        ) from exc
    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n\n".join(parts)


EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
}
